#!/usr/bin/env python3
"""
Translate a book's DOCX to all site languages and export per-language PDFs.

Two books are supported:
  carrom-techniques-and-skills   (default) — Arun's coaching book
  carrom-players-guide                     — Players Guide (English source)

Reads:  static/downloads/{SourceFilename}.docx
Writes: static/downloads/{book-slug}-{lang}.pdf
Cache:  .cache/book-docx/{book-slug}-{lang}.docx

Requires: python-docx, deep-translator, LibreOffice (soffice on PATH).

Usage:
  python3 scripts/generate-book-pdfs.py                              # techniques, all languages
  python3 scripts/generate-book-pdfs.py da de                        # techniques, specific languages
  python3 scripts/generate-book-pdfs.py --book carrom-players-guide  # players-guide, all languages
  python3 scripts/generate-book-pdfs.py --book carrom-players-guide en da  # players-guide, EN + DA
  python3 scripts/generate-book-pdfs.py --force da                   # re-translate even if cached
  python3 scripts/generate-book-pdfs.py --pdf-only                   # convert cached docx only
"""

from __future__ import annotations

import importlib.util
import re
import shutil
import subprocess
import sys
import time
from pathlib import Path

from docx import Document
from deep_translator import GoogleTranslator

REPO = Path(__file__).resolve().parents[1]
CACHE_DIR = REPO / ".cache" / "book-docx"
OUTPUT_DIR = REPO / "static" / "downloads"

# Registry of books this pipeline can translate. Add a new entry to
# translate a new book — no other code change needed.
BOOKS = {
    "carrom-techniques-and-skills": REPO / "static" / "downloads" / "CarromTechniqandSkills.docx",
    "carrom-players-guide": REPO / "static" / "downloads" / "PlayersGuideEnglish.docx",
}
DEFAULT_BOOK = "carrom-techniques-and-skills"

ALL_LANGS = ["en", "da", "de", "mr", "it", "fr", "si", "hi", "gu", "pl", "mni", "ta", "te", "or", "bn", "as", "cs", "sr", "sv", "bg", "ur"]

# Reuse translation helpers from translate-books.py
_spec = importlib.util.spec_from_file_location(
    "translate_books", REPO / "scripts" / "translate-books.py"
)
_tb = importlib.util.module_from_spec(_spec)
_spec.loader.exec_module(_tb)
translate_text = _tb.translate_text
get_translator = _tb.get_translator


def find_soffice() -> str:
    for candidate in (
        shutil.which("soffice"),
        "/opt/homebrew/bin/soffice",
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
    ):
        if candidate and Path(candidate).exists():
            return candidate
    raise RuntimeError(
        "LibreOffice not found. Install with: brew install --cask libreoffice"
    )


def docx_path(book_slug: str, lang: str) -> Path:
    return CACHE_DIR / f"{book_slug}-{lang}.docx"


def pdf_path(book_slug: str, lang: str) -> Path:
    return OUTPUT_DIR / f"{book_slug}-{lang}.pdf"


def iter_paragraphs(doc: Document):
    for para in doc.paragraphs:
        yield para
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    yield para
    for section in doc.sections:
        for part in (section.header, section.footer):
            if part is None:
                continue
            for para in part.paragraphs:
                yield para
            for table in part.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            yield para


def should_translate(text: str) -> bool:
    text = text.strip()
    if not text:
        return False
    if re.fullmatch(r"[\d\s.,/\-–—%:;]+", text):
        return False
    return bool(re.search(r"[A-Za-z]", text))


def set_paragraph_text(para, text: str) -> None:
    if para.runs:
        para.runs[0].text = text
        for run in para.runs[1:]:
            run.text = ""
    else:
        para.add_run(text)


# Font overrides for scripts whose characters Calibri can't render.
SCRIPT_FONTS = {
    "mni": "Noto Sans Meetei Mayek",
}


def apply_script_font(doc: Document, font_name: str) -> None:
    """Set every run that contains non-ASCII text to the given font.

    Sets all four font slots (ascii, hAnsi, cs, eastAsia) so Word/LibreOffice
    uses the override regardless of which script slot it considers a glyph
    to belong to.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    for para in iter_paragraphs(doc):
        for run in para.runs:
            if not run.text or not any(ord(c) > 127 for c in run.text):
                continue
            rPr = run._element.get_or_add_rPr()
            # Remove any existing rFonts
            for existing in rPr.findall(qn("w:rFonts")):
                rPr.remove(existing)
            rFonts = OxmlElement("w:rFonts")
            for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
                rFonts.set(qn(attr), font_name)
            rPr.insert(0, rFonts)


def translate_docx(src: Path, dest: Path, lang: str) -> None:
    print(f"  translating docx → {lang} …", flush=True)
    doc = Document(str(src))
    translator = get_translator(lang)
    total = sum(1 for _ in iter_paragraphs(doc))
    done = 0

    for para in iter_paragraphs(doc):
        done += 1
        raw = para.text
        if not should_translate(raw):
            continue
        translated = translate_text(raw, translator)
        if translated != raw:
            set_paragraph_text(para, translated)
        if done % 40 == 0:
            print(f"    … {done}/{total} paragraphs", flush=True)
        time.sleep(0.05)

    if lang in SCRIPT_FONTS:
        print(f"  applying font override: {SCRIPT_FONTS[lang]}", flush=True)
        apply_script_font(doc, SCRIPT_FONTS[lang])

    dest.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dest))
    print(f"  wrote {dest.relative_to(REPO)}", flush=True)


def copy_en_docx(source_docx: Path, dest: Path) -> None:
    dest.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(source_docx, dest)
    print(f"  copied EN source → {dest.relative_to(REPO)}", flush=True)


def convert_to_pdf(docx: Path, pdf: Path) -> None:
    soffice = find_soffice()
    pdf.parent.mkdir(parents=True, exist_ok=True)
    outdir = pdf.parent
    cmd = [
        soffice,
        "--headless",
        "--convert-to",
        "pdf",
        "--outdir",
        str(outdir),
        str(docx),
    ]
    print(f"  converting PDF …", flush=True)
    subprocess.run(cmd, check=True, capture_output=True, text=True)
    generated = outdir / f"{docx.stem}.pdf"
    if not generated.exists():
        raise RuntimeError(f"LibreOffice did not produce {generated}")
    if generated != pdf:
        generated.replace(pdf)
    size_mb = pdf.stat().st_size / (1024 * 1024)
    print(f"  wrote {pdf.relative_to(REPO)} ({size_mb:.1f} MB)", flush=True)


def process_lang(book_slug: str, source_docx: Path, lang: str, *, force: bool, pdf_only: bool) -> None:
    print(f"\n=== {book_slug} — {lang.upper()} ===", flush=True)
    if not source_docx.exists():
        raise FileNotFoundError(f"Source docx missing: {source_docx}")

    cached = docx_path(book_slug, lang)
    out_pdf = pdf_path(book_slug, lang)

    if not pdf_only:
        if lang == "en":
            if force or not cached.exists():
                copy_en_docx(source_docx, cached)
        elif force or not cached.exists():
            translate_docx(source_docx, cached, lang)
        elif not cached.exists():
            raise FileNotFoundError(f"No cached docx for {lang}; run without --pdf-only")

    if not cached.exists():
        raise FileNotFoundError(f"Missing docx for {lang}: {cached}")

    convert_to_pdf(cached, out_pdf)


def main() -> None:
    # Parse argv: --force, --pdf-only, --book <slug>, and positional lang codes.
    raw = sys.argv[1:]
    force = "--force" in raw
    pdf_only = "--pdf-only" in raw

    book_slug = DEFAULT_BOOK
    positional: list[str] = []
    i = 0
    while i < len(raw):
        tok = raw[i]
        if tok == "--book":
            if i + 1 >= len(raw):
                print("--book requires a slug argument", file=sys.stderr)
                sys.exit(1)
            book_slug = raw[i + 1]
            i += 2
            continue
        if tok in ("--force", "--pdf-only"):
            i += 1
            continue
        if tok.startswith("-"):
            print(f"Unknown flag: {tok}", file=sys.stderr)
            sys.exit(1)
        positional.append(tok)
        i += 1

    if book_slug not in BOOKS:
        print(f"Unknown book: {book_slug}. Known: {', '.join(BOOKS)}", file=sys.stderr)
        sys.exit(1)
    source_docx = BOOKS[book_slug]

    targets = positional or ALL_LANGS
    for lang in targets:
        if lang not in ALL_LANGS:
            print(f"Unknown language: {lang}", file=sys.stderr)
            sys.exit(1)

    for lang in targets:
        process_lang(book_slug, source_docx, lang, force=force, pdf_only=pdf_only)

    print("\nDone.", flush=True)


if __name__ == "__main__":
    main()
